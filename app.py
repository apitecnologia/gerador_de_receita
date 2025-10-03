import os
import io
import pandas as pd
from flask import Flask, render_template, request, redirect, url_for, flash, send_file
from flask_sqlalchemy import SQLAlchemy
from flask_login import LoginManager, UserMixin, login_user, logout_user, login_required, current_user
from werkzeug.security import generate_password_hash, check_password_hash
from docx import Document
from functools import wraps
import getpass
from flask_migrate import Migrate

# 1. CONFIGURAÇÃO DA APLICAÇÃO
app = Flask(__name__)

# Garante que o diretório 'instance' exista ANTES de qualquer configuração
instance_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'instance')
os.makedirs(instance_path, exist_ok=True)

# Configurações para produção (Render) e desenvolvimento (local)
app.config['SECRET_KEY'] = os.environ.get('SECRET_KEY', os.urandom(24))
app.config['SQLALCHEMY_DATABASE_URI'] = os.environ.get('DATABASE_URL', f'sqlite:///{os.path.join(instance_path, "database.sqlite")}')
app.config['SQLALCHEMY_TRACK_MODIFICATIONS'] = False

# Inicializa as extensões
db = SQLAlchemy(app)
migrate = Migrate(app, db)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message = "Por favor, faça o login para acessar esta página."
login_manager.login_message_category = "error"


# 2. MODELOS DO BANCO DE DADOS
class User(UserMixin, db.Model):
    id = db.Column(db.Integer, primary_key=True)
    username = db.Column(db.String(80), unique=True, nullable=False)
    password_hash = db.Column(db.String(256))
    full_name = db.Column(db.String(150))
    phone = db.Column(db.String(20))
    email = db.Column(db.String(120))
    is_admin = db.Column(db.Boolean, default=False, nullable=False)

    def set_password(self, password):
        self.password_hash = generate_password_hash(password)

    def check_password(self, password):
        return check_password_hash(self.password_hash, password)

class Product(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(500), nullable=False) # Limite de caracteres aumentado
    concentration = db.Column(db.String(100))
    unit = db.Column(db.String(50))
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=False)

    user = db.relationship('User', backref=db.backref('products', lazy=True, cascade="all, delete-orphan"))

# Função para carregar o usuário da sessão
@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))

# Decorator para proteger rotas de administrador
def admin_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if not current_user.is_authenticated or not current_user.is_admin:
            flash('Acesso negado. Esta área é apenas para administradores.', 'error')
            return redirect(url_for('index'))
        return f(*args, **kwargs)
    return decorated_function


# 3. ROTAS DA APLICAÇÃO
@app.route('/login', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        user = User.query.filter_by(username=username).first()
        if not user or not user.check_password(password):
            flash('Usuário ou senha inválidos. Por favor, tente novamente.', 'error')
            return redirect(url_for('login'))
        login_user(user)
        return redirect(url_for('index'))
    return render_template('login.html')

@app.route('/register', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')
        user = User.query.filter_by(username=username).first()
        if user:
            flash('Este nome de usuário já existe.', 'error')
            return redirect(url_for('register'))
        new_user = User(username=username)
        new_user.set_password(password)
        db.session.add(new_user)
        db.session.commit()
        flash('Cadastro realizado com sucesso! Faça o login.', 'success')
        return redirect(url_for('login'))
    return render_template('register.html')

@app.route('/logout')
@login_required
def logout():
    logout_user()
    return redirect(url_for('login'))

@app.route('/')
@login_required
def index():
    products = Product.query.filter_by(user_id=current_user.id).order_by(Product.name).all()
    if not products:
        flash('Nenhum produto encontrado. Por favor, importe uma planilha para começar.', 'success')
        return redirect(url_for('upload_file'))
    return render_template('index.html', produtos=products)

@app.route('/upload', methods=['GET', 'POST'])
@login_required
def upload_file():
    if request.method == 'POST':
        if 'file' not in request.files or request.files['file'].filename == '':
            flash('Nenhum arquivo selecionado.', 'error')
            return redirect(request.url)
        file = request.files['file']
        if file and file.filename.endswith('.xlsx'):
            try:
                df = pd.read_excel(file)
                colunas_esperadas = ['PRODUTO', 'CONCENTRACAO', 'UNIDADE']
                if not all(col in df.columns for col in colunas_esperadas):
                    flash("Erro: A planilha deve conter as colunas: PRODUTO, CONCENTRACAO, UNIDADE.", 'error')
                    return redirect(request.url)
                Product.query.filter_by(user_id=current_user.id).delete()
                for _, row in df.iterrows():
                    new_product = Product(
                        name=str(row['PRODUTO']).strip(),
                        concentration=str(row['CONCENTRACAO']),
                        unit=str(row['UNIDADE']),
                        user_id=current_user.id
                    )
                    db.session.add(new_product)
                db.session.commit()
                flash('Planilha importada com sucesso!', 'success')
                return redirect(url_for('index'))
            except Exception as e:
                db.session.rollback()
                flash(f'Erro ao processar a planilha: {e}', 'error')
                return redirect(request.url)
        else:
            flash('Formato de arquivo inválido. Apenas .xlsx é permitido.', 'error')
            return redirect(request.url)
    return render_template('upload.html')

@app.route('/limpar_dados')
@login_required
def limpar_dados():
    try:
        Product.query.filter_by(user_id=current_user.id).delete()
        db.session.commit()
        flash('Sua lista de produtos foi limpa com sucesso.', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Ocorreu um erro ao limpar os dados: {e}', 'error')
    return redirect(url_for('upload_file'))

@app.route('/gerar_receita', methods=['POST'])
@login_required
def gerar_receita():
    ids_selecionados = request.form.getlist('produtos')
    if not ids_selecionados:
        flash('Nenhum produto foi selecionado para gerar a receita.', 'error')
        return redirect(url_for('index'))
    doc = Document()
    doc.add_heading('Prescrição', 0)
    doc.add_paragraph(f"Data: {pd.Timestamp.now().strftime('%d/%m/%Y')}")
    doc.add_paragraph('\nPrezado cliente, segue a relação de produtos para sua aprovação:\n')
    for produto_id in ids_selecionados:
        produto = Product.query.get(int(produto_id))
        if produto and produto.user_id == current_user.id:
            quantidade = request.form.get(f"quantidade_{produto_id}", 1)
            texto_produto = f"{produto.name} - {produto.concentration} - {quantidade} {produto.unit}"
            doc.add_paragraph(texto_produto, style='List Bullet')
    doc.add_paragraph('\n')
    doc.add_paragraph('Atenciosamente,')
    doc.add_paragraph('Victor Monteiro - Executivo de Vendas')
    doc.add_paragraph('Tel: (11) 96712-3704')
    doc.add_paragraph('Email: comercial.vmbusiness@gmail.com')
    doc_buffer = io.BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return send_file(
        doc_buffer,
        as_attachment=True,
        download_name='receita_comercial.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# ... (Rotas de Admin permanecem as mesmas)

@app.cli.command("create-admin")
def create_admin():
    """Cria um usuário administrador inicial."""
    # ... (Código do create-admin permanece o mesmo)

if __name__ == '__main__':
    app.run(debug=True)